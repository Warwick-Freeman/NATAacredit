"""
convert_policies.py
Converts AASM Policy/Protocol/Form .docx files to styled HTML documents
for the Nexus 360 Accreditation system.
"""

import sys
import os
import re
import html as html_module
from pathlib import Path

try:
    import docx
except ImportError:
    print("python-docx not found. Installing...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    import docx

# ---------------------------------------------------------------------------
# Document ID mapping: (relative_path_fragment, doc_id, aasm_standard, title)
# relative_path_fragment is matched against the full docx path
# ---------------------------------------------------------------------------
MAPPING = [
    # Network
    (r"1 - Network.*Continuing Education", "POL-NET-001", "AASM N-5, N-7, N-10", "Continuing Education Policy"),
    (r"1 - Network.*Employee Background Check", "POL-NET-002", "AASM N-15", "Employee Background Check Policy"),
    (r"1 - Network.*Sleep Technicians", "POL-NET-003", "AASM N-8, N-9, N-11", "Sleep Technicians and Technologists Policy"),
    (r"1 - Network.*Direct Referral", "POL-NET-004", "AASM N-19", "Direct Referral Policy"),
    (r"1 - Network.*Maintenance and Organization", "POL-NET-005", "AASM N-18", "Maintenance and Organization of Medical Records Policy"),
    (r"1 - Network.*Emergency Policy", "POL-NET-006", "AASM N-20", "Emergency Policy"),
    (r"1 - Network.*Quality Assurance", "POL-NET-007", "AASM N-22, N-23", "Quality Assurance Policy"),
    (r"Inter-Scorer Reliability Assessment - Not Using AASM", "POL-NET-008", "AASM N-24", "Inter-Scorer Reliability Assessment — Not Using AASM Sleep ISR"),
    (r"Inter-Scorer Reliability Assessment - Use of AASM", "POL-NET-009", "AASM N-24", "Inter-Scorer Reliability Assessment — Use of AASM Sleep ISR"),
    (r"1 - Network.*Significant Adverse Events", "POL-NET-010", "AASM N-21", "Significant Adverse Events Policy"),
    # Clinic
    (r"2 - Clinic.*Sleep Clinic Patient Acceptance", "POL-CLI-001", "AASM C-1, S-5", "Sleep Clinic Patient Acceptance Policy"),
    (r"2 - Clinic.*Sleep Clinic Patient Management", "POL-CLI-002", "AASM C-1, C-2", "Sleep Clinic Patient Management and Assessment Policy"),
    (r"2 - Clinic.*Medical CPR AED Emergency Drill", "POL-CLI-003", "AASM S-6, S-7", "Medical CPR/AED Emergency Drill Policy"),
    (r"2 - Clinic.*Hazardous Materials", "POL-CLI-004", "AASM S-10", "Hazardous Materials Management Policy"),
    (r"2 - Clinic.*Occupational Safety", "POL-CLI-005", "AASM S-9", "Occupational Safety Management Policy"),
    (r"2 - Clinic.*Patient Safety Risk Analysis", "POL-CLI-006", "AASM S-11", "Patient Safety Risk Analysis Policy"),
    (r"2 - Clinic.*Significant Adv[ae]re", "POL-CLI-007", "AASM S-12, N-21", "Significant Adverse Events Policy"),
    # Diagnostics/Lab — Policies
    (r"3 - Diagnostics.*In-Lab Patient Acceptance", "POL-LAB-001", "AASM L-1, S-5", "In-Lab Patient Acceptance Policy"),
    (r"3 - Diagnostics.*In-Lab Emergency Plan", "POL-LAB-002", "AASM L-23", "In-Lab Emergency Plan Policy"),
    (r"3 - Diagnostics.*Medical CPR AED Emergency Drill", "POL-LAB-003", "AASM S-6, S-7", "Medical CPR/AED Emergency Drill Policy"),
    (r"3 - Diagnostics.*In-Lab Equipment Maintenance", "POL-LAB-004", "AASM L-15", "In-Lab Equipment Maintenance Policy"),
    (r"3 - Diagnostics.*Hazardous Materials", "POL-LAB-005", "AASM S-10", "Hazardous Materials Management Policy"),
    (r"3 - Diagnostics.*In-Lab Safety Risk", "POL-LAB-006", "AASM L-24", "In-Lab Safety Risk Prevention Policy"),
    (r"3 - Diagnostics.*Occupational Safety", "POL-LAB-007", "AASM S-9", "Occupational Safety Management Policy"),
    (r"3 - Diagnostics.*Patient Safety Risk Analysis", "POL-LAB-008", "AASM S-11", "Patient Safety Risk Analysis Policy"),
    (r"3 - Diagnostics.*4\. Safety.*Significant Adverse", "POL-LAB-009", "AASM S-12", "Significant Adverse Event Reporting Policy"),
    # Diagnostics/Lab — Protocols
    (r"3 - Diagnostics.*BiPAP", "PRO-LAB-001", "AASM L-14, L-16", "BiPAP Protocol"),
    (r"3 - Diagnostics.*CO2 Monitoring", "PRO-LAB-002", "AASM L-17", "CO2 Monitoring Protocol"),
    (r"3 - Diagnostics.*CPAP", "PRO-LAB-003", "AASM L-14, L-16", "CPAP Protocol"),
    (r"3 - Diagnostics.*MSLT", "PRO-LAB-004", "AASM L-16", "MSLT Protocol"),
    (r"3 - Diagnostics.*MWT", "PRO-LAB-005", "AASM L-16", "MWT Protocol"),
    (r"3 - Diagnostics.*Neonate", "PRO-LAB-006", "AASM L-18", "Neonate and Infant PSG Protocol"),
    (r"3 - Diagnostics.*Pediatric PAP Titration", "PRO-LAB-007", "AASM L-18", "Pediatric PAP Titration Protocol"),
    (r"3 - Diagnostics.*Pediatric PSG", "PRO-LAB-008", "AASM L-18", "Pediatric PSG Protocol"),
    (r"3 - Diagnostics.*Polysomnography", "PRO-LAB-009", "AASM L-16, L-19, L-20", "Adult Polysomnography Protocol"),
    (r"3 - Diagnostics.*Split Night", "PRO-LAB-010", "AASM L-16", "Split Night Protocol"),
    # HSAT — Policies
    (r"4 - Home Sleep.*HSAT Patient Acceptance", "POL-HST-001", "AASM H-1, S-5", "HSAT Patient Acceptance Policy"),
    (r"4 - Home Sleep.*Medical CPR AED Emergency Drill", "POL-HST-002", "AASM S-6, S-7", "Medical CPR/AED Emergency Drill Policy"),
    (r"4 - Home Sleep.*HSAT Training", "POL-HST-003", "AASM H-3", "HSAT Training Policy"),
    (r"4 - Home Sleep.*Hazardous Materials", "POL-HST-004", "AASM S-10", "Hazardous Materials Management Policy"),
    (r"4 - Home Sleep.*Occupational.*Safety", "POL-HST-005", "AASM S-9", "Occupational Safety Management Policy"),
    (r"4 - Home Sleep.*Patient Safety Risk", "POL-HST-006", "AASM S-11", "Patient Safety Risk Analysis Policy"),
    (r"4 - Home Sleep.*Significant Adverse", "POL-HST-007", "AASM S-12", "Significant Adverse Event Reporting Policy"),
    # HSAT — Protocol
    (r"4 - Home Sleep.*HSAT Protocol", "PRO-HST-001", "AASM H-4, H-5, H-6", "HSAT Protocol"),
    # DME
    (r"5 - DME.*DME Patient Acceptance", "POL-DME-001", "AASM D-1, D-2", "DME Patient Acceptance, Ordering and Equipment Policy"),
    (r"5 - DME.*DME Training and Education", "POL-DME-002", "AASM D-3, D-4", "DME Training and Education Policy"),
    (r"5 - DME.*DME Equipment Maintenance", "POL-DME-003", "AASM D-5", "DME Equipment Maintenance Policy"),
    (r"5 - DME.*Billing Discrepancy", "POL-DME-004", "AASM D-6", "Billing Discrepancy Resolution Policy"),
    (r"5 - DME.*Equipment Loaner", "POL-DME-005", "AASM D-6", "Equipment Loaner Policy"),
    (r"5 - DME.*Financial Hardship", "POL-DME-006", "AASM D-6", "Financial Hardship / Charity Care Policy"),
    (r"5 - DME.*Fraud, Waste", "POL-DME-007", "AASM D-6", "Fraud, Waste and Abuse Policy"),
    (r"5 - DME.*Ongoing Patient Management", "POL-DME-008", "AASM D-10", "Ongoing Patient Management Policy"),
    (r"5 - DME.*6 .Safety.*Hazardous", "POL-DME-009", "AASM S-10", "Hazardous Materials Management Policy"),
    (r"5 - DME.*OSHA", "POL-DME-010", "AASM S-9", "OSHA Occupational Safety Policy"),
    (r"5 - DME.*6 .Safety.*Patient Safety Risk", "POL-DME-011", "AASM S-11", "Patient Safety Risk Analysis Policy"),
    (r"5 - DME.*Significant Adverse Event Reporting", "POL-DME-012", "AASM S-12", "Significant Adverse Event Reporting Policy"),
    (r"5 - DME.*Medical Emergency CPR", "POL-DME-013", "AASM S-6, S-7", "Medical Emergency CPR/AED Drill Policy"),
    # Forms
    (r"6 - Forms.*Bomb Threat", "FRM-AASM-001", "", "Bomb Threat Form"),
    (r"6 - Forms.*Cardiopulmonary Emergency Drill", "FRM-AASM-002", "", "Cardiopulmonary Emergency Drill Form"),
    (r"6 - Forms.*CE Log.*Inservice", "FRM-AASM-003", "", "CE Log — Inservice Sign-In Sheet"),
    (r"6 - Forms.*CEC Log.*Technical", "FRM-AASM-004", "", "CEC Log — Technical Staff"),
    (r"6 - Forms.*CME Log.*Professional", "FRM-AASM-005", "", "CME Log — Professional Staff"),
    (r"6 - Forms.*Direct Referral Form", "FRM-AASM-006", "", "Direct Referral Form"),
    (r"6 - Forms.*DME Equipment Maintenance Failure", "FRM-AASM-007", "", "DME Equipment Maintenance Failure Log"),
    (r"6 - Forms.*Fax Request", "FRM-AASM-008", "", "Fax Request Form"),
    (r"6 - Forms.*HSAT Equipment Maintenance Failure", "FRM-AASM-009", "", "HSAT Equipment Maintenance Failure Log"),
    (r"6 - Forms.*In-Lab Equipment Maintenance Failure", "FRM-AASM-010", "", "In-Lab Equipment Maintenance Failure Log"),
    (r"6 - Forms.*Monthly Visual Inspection", "FRM-AASM-011", "", "Monthly Visual Inspection Log"),
    (r"6 - Forms.*Patient Safety Risk Analysis Checklist", "FRM-AASM-012", "", "Patient Safety Risk Analysis Checklist"),
    (r"6 - Forms.*QA Sample Report", "FRM-AASM-013", "", "QA Sample Report"),
    (r"6 - Forms.*Significant Adverse Event Form", "FRM-AASM-014", "", "Significant Adverse Event Form"),
    (r"6 - Forms.*Sub-contract Review", "FRM-AASM-015", "", "Sub-contract Review Sheet"),
]

# ---------------------------------------------------------------------------
# CSS / HTML template
# ---------------------------------------------------------------------------
CSS = """
@page { size: A4; margin: 18mm 16mm 22mm 16mm; }
body { font-family: "Segoe UI", Arial, sans-serif; font-size: 11pt; color: #1f2933; line-height: 1.45; margin: 0; padding: 24px; background: #f5f7fa; }
.page { max-width: 920px; margin: 0 auto; background: #ffffff; padding: 32px 40px; border: 1px solid #d6dde6; border-radius: 6px; box-shadow: 0 2px 6px rgba(0,0,0,0.04); }
.doc-header { display: grid; grid-template-columns: 1fr 1fr; gap: 8px 18px; border-bottom: 3px solid #14406b; padding-bottom: 10px; margin-bottom: 18px; }
.doc-header .kind { grid-column: 1 / span 2; font-size: 9.5pt; color: #506478; text-transform: uppercase; letter-spacing: 1.5px; }
.doc-header .title { grid-column: 1 / span 2; font-size: 18pt; font-weight: 700; color: #14406b; margin: 0 0 4px 0; }
.doc-header .subtitle { grid-column: 1 / span 2; font-size: 10.5pt; color: #506478; margin-bottom: 6px; }
.doc-header .field { font-size: 9.5pt; color: #1f2933; border-bottom: 1px dotted #ccd4dc; padding-bottom: 3px; }
.doc-header .field strong { display: inline-block; width: 130px; color: #506478; font-weight: 600; }
h2 { font-size: 12.5pt; color: #14406b; margin-top: 22px; margin-bottom: 8px; border-left: 4px solid #14406b; padding: 2px 0 2px 8px; background: #eef3f9; }
h3 { font-size: 11pt; color: #1f2933; margin-top: 16px; margin-bottom: 4px; }
p, li { font-size: 10.5pt; }
ul, ol { margin: 4px 0 8px 22px; padding: 0; }
table { width: 100%; border-collapse: collapse; margin: 8px 0 12px 0; font-size: 10pt; }
th, td { border: 1px solid #c5cfdb; padding: 6px 8px; text-align: left; vertical-align: top; }
th { background: #eef3f9; color: #14406b; font-weight: 600; }
.callout { background: #fff7e0; border: 1px solid #f0d27e; border-left: 4px solid #d49b1c; padding: 8px 12px; margin: 10px 0; font-size: 10pt; }
.signoff { display: grid; grid-template-columns: 1fr 1fr 1fr; gap: 12px; margin-top: 14px; }
.sign-cell { border: 1px solid #c5cfdb; padding: 8px; min-height: 60px; font-size: 9.5pt; color: #506478; }
.sign-cell strong { color: #1f2933; }
.footer { border-top: 1px solid #c5cfdb; margin-top: 24px; padding-top: 8px; font-size: 9pt; color: #506478; text-align: center; }
.clause-tag { display: inline-block; background: #14406b; color: #fff; font-size: 8.5pt; padding: 1px 6px; border-radius: 3px; margin-left: 4px; vertical-align: middle; }
.fill { display: inline-block; min-width: 180px; border-bottom: 1px solid #1f2933; height: 14px; }
.fill-sm { display: inline-block; min-width: 80px; border-bottom: 1px solid #1f2933; height: 14px; }
.box { display: inline-block; width: 12px; height: 12px; border: 1px solid #1f2933; vertical-align: middle; margin-right: 4px; }
@media print { body { background: #ffffff; padding: 0; } .page { box-shadow: none; border: none; padding: 0; } }
""".strip()

# ---------------------------------------------------------------------------
# Text substitution helpers
# ---------------------------------------------------------------------------

SUBSTITUTIONS = [
    # Remove SAMPLE headers
    (re.compile(r'^\s*S+AMPLE\s*(POLICY|Protocol|PROTOCOL)?\s*$', re.IGNORECASE), ""),
    # Logbook / equipment log
    (re.compile(r'\b(equipment\s+log(?:book)?|maintenance\s+log(?:book)?|logbook|calibration\s+log|maintenance\s+record)\b', re.IGNORECASE),
     "the Nexus 360 Equipment Register"),
    # Personnel file
    (re.compile(r'\b(personnel\s+file|personnel\s+folder|staff\s+file)\b', re.IGNORECASE),
     "the Nexus 360 Staff &amp; Training module"),
    # (insert location)
    (re.compile(r'\(?\[?\s*insert\s+location\s*\]?\)?', re.IGNORECASE), "Nexus 360"),
    # "on file" when referring to documents
    (re.compile(r'\bon file\b', re.IGNORECASE), "in Nexus 360"),
    # monthly staff meeting
    (re.compile(r'\bmonthly\s+staff\s+meeting\b', re.IGNORECASE),
     "monthly staff meeting (documented in Nexus 360)"),
    # equipment failure / NC reporting
    (re.compile(r'\breport(?:ed|ing)?\s+(?:equipment\s+)?(?:failure|malfunction)s?\s+(?:to|via|using)?\s*(?:the\s+)?(?:NC\s*&?\s*CAPA|nonconformance)\s+(?:module|system|register)?\b', re.IGNORECASE),
     "reported via the Nexus 360 NC &amp; CAPA module"),
    # QA report storage
    (re.compile(r'\bQA\s+(?:report|summary)\s+(?:is\s+)?(?:stored|filed|kept|saved)\b', re.IGNORECASE),
     "QA report stored in the Nexus 360 Quality Indicators module"),
    # "To be signed by the network director" / "site director"
    (re.compile(r'To\s+be\s+signed\s+by\s+the\s+(?:network|site)\s+director', re.IGNORECASE),
     '<span class="fill"></span>'),
    # phone placeholders  (___) ___-____
    (re.compile(r'\(___\)\s*___-____'), '<span class="fill-sm"></span>'),
    # 00/00/0000
    (re.compile(r'\b0{1,2}/0{1,2}/0{2,4}\b'), '<span class="fill-sm"></span>'),
    # [enter date] / (enter date)
    (re.compile(r'\[?enter\s+date\]?', re.IGNORECASE), '<span class="fill-sm"></span>'),
    # Checkbox character □ → span.box
    (re.compile(r'□'), '<span class="box"></span>'),
    # Blank line fillers (many tabs)
    (re.compile(r'(\t{5,})'), lambda m: '<span class="fill"></span> ' * (len(m.group(1)) // 5)),
    # Inline fill: trailing tabs after a label
    (re.compile(r'(\t{2,})'), ' <span class="fill"></span> '),
]


def apply_substitutions(text: str) -> str:
    for pattern, replacement in SUBSTITUTIONS:
        if callable(replacement):
            text = pattern.sub(replacement, text)
        else:
            text = pattern.sub(replacement, text)
    return text


def is_section_header(style_name: str, text: str) -> bool:
    """Returns True if this paragraph is a section heading (PURPOSE, POLICY, PROCEDURE, etc.)"""
    keywords = {"PURPOSE", "POLICY", "PROCEDURE", "BACKGROUND", "SCOPE",
                "RESPONSIBILITIES", "REFERENCES", "DEFINITIONS", "OVERVIEW",
                "GENERAL", "PROTOCOL", "INSTRUCTIONS", "INTRODUCTION"}
    if "H4" in style_name or "Sample Body Subhead" in style_name:
        return True
    if text.strip().upper() in keywords:
        return True
    return False


def is_doc_title(style_name: str) -> bool:
    return "H3" in style_name and "Section Subhead" in style_name


def is_meta_line(text: str) -> bool:
    """Skip metadata lines like Effective Date, Revision Date, Approval Signature."""
    stripped = text.strip()
    if re.match(r'^(Effective Date|Revision Date|Approval Signature)', stripped, re.IGNORECASE):
        return True
    return False


def is_sample_header(style_name: str, text: str) -> bool:
    """Detect SAMPLE POLICY / SSAMPLE headers."""
    if "H5" in style_name and "Sample" in style_name:
        return True
    if re.match(r'^\s*S+AMPLE\s*(POLICY|Protocol)?\s*$', text, re.IGNORECASE):
        return True
    return False


def classify_list_style(style_name: str):
    """Returns (is_list, tier, is_ordered)"""
    if "Sample Num List" in style_name or "Num List" in style_name:
        if "Tier 1" in style_name:
            return True, 1, True
        elif "Tier 2" in style_name:
            return True, 2, True
        elif "Tier 3" in style_name:
            return True, 3, True
        return True, 1, True
    if "List Bullet" in style_name or "Bullet" in style_name:
        return True, 1, False
    if "List Number" in style_name:
        return True, 1, True
    return False, 0, False


SECTION_NAMES = {
    "PURPOSE": "Purpose",
    "POLICY": "Policy",
    "PROCEDURE": "Procedure",
    "PROCEDURES": "Procedure",
    "PROTOCOL": "Protocol",
    "BACKGROUND": "Background",
    "SCOPE": "Scope",
    "RESPONSIBILITIES": "Responsibilities",
    "REFERENCES": "References",
    "DEFINITIONS": "Definitions",
    "OVERVIEW": "Overview",
    "GENERAL": "General",
    "INSTRUCTIONS": "Instructions",
    "INTRODUCTION": "Introduction",
}


def get_section_number_and_name(text: str, section_counter: list) -> str:
    """Return numbered section heading HTML."""
    upper = text.strip().upper()
    nice = SECTION_NAMES.get(upper, text.strip().title())
    section_counter[0] += 1
    return f"<h2>{section_counter[0]}. {nice}</h2>"


def infer_authorised_by(doc_id: str) -> str:
    if "NET" in doc_id:
        return "Network Director"
    elif "CLI" in doc_id:
        return "Clinic Medical Director"
    elif "LAB" in doc_id:
        return "Laboratory Director"
    elif "HST" in doc_id:
        return "HSAT Program Director"
    elif "DME" in doc_id:
        return "DME Program Director"
    else:
        return "Program Director"


def render_table(table) -> str:
    rows = table.rows
    if not rows:
        return ""
    html_rows = []
    # Use first row as header if it looks like one (all cells have text)
    first_row_texts = [c.text.strip() for c in rows[0].cells]
    use_header = all(t for t in first_row_texts)

    html_rows.append("<table>")
    for i, row in enumerate(rows):
        cells = row.cells
        if i == 0 and use_header:
            html_rows.append("<tr>")
            for cell in cells:
                txt = apply_substitutions(html_module.escape(cell.text.strip()))
                html_rows.append(f"<th>{txt}</th>")
            html_rows.append("</tr>")
        else:
            html_rows.append("<tr>")
            for cell in cells:
                raw = cell.text.strip()
                # Convert cell text: newlines → <br>, apply subs
                raw_esc = apply_substitutions(html_module.escape(raw))
                raw_esc = raw_esc.replace("\n", "<br>")
                html_rows.append(f"<td>{raw_esc}</td>")
            html_rows.append("</tr>")
    html_rows.append("</table>")
    return "\n".join(html_rows)


def paragraphs_to_html(paragraphs, tables, is_form: bool) -> str:
    """Convert a list of paragraphs (and tables) into HTML body content."""
    lines = []
    section_counter = [0]
    list_stack = []  # stack of (tier, is_ordered)

    # Build an index of table anchors — docx stores tables inline in the body XML.
    # We'll interleave them by tracking a table index.
    # For simplicity, we append all tables at the end if there is no better anchor.
    # To interleave correctly we use the paragraph _p element positions.
    # This is a simplified approach: tables are emitted after all paragraphs
    # unless we detect them by position.

    def close_lists_to(target_tier: int):
        while list_stack and list_stack[-1][0] > target_tier:
            tier, ordered = list_stack.pop()
            tag = "ol" if ordered else "ul"
            lines.append(f"</{tag}>")

    def close_all_lists():
        while list_stack:
            tier, ordered = list_stack.pop()
            tag = "ol" if ordered else "ul"
            lines.append(f"</{tag}>")

    for para in paragraphs:
        style = para.style.name
        text = para.text

        # Skip SAMPLE headers
        if is_sample_header(style, text):
            continue

        # Skip doc-title line (already in header)
        if is_doc_title(style):
            continue

        # Skip metadata lines
        if is_meta_line(text):
            continue

        # Skip blank paragraphs (but flush lists first if needed)
        if not text.strip():
            close_all_lists()
            continue

        # Section headers
        if is_section_header(style, text):
            close_all_lists()
            lines.append(get_section_number_and_name(text, section_counter))
            continue

        # List items
        is_list, tier, is_ordered = classify_list_style(style)
        if is_list:
            # Manage list nesting
            close_lists_to(tier)
            if not list_stack or list_stack[-1][0] < tier:
                tag = "ol" if is_ordered else "ul"
                lines.append(f"<{tag}>")
                list_stack.append((tier, is_ordered))
            txt = apply_substitutions(html_module.escape(text.strip()))
            lines.append(f"<li>{txt}</li>")
            continue

        # Regular body paragraph
        close_all_lists()
        txt = apply_substitutions(html_module.escape(text.strip()))

        # Detect sub-headings (bold-only short lines)
        if (("No Spacing" not in style) and len(text.strip()) < 80 and
                any(run.bold for run in para.runs if run.text.strip())):
            lines.append(f"<h3>{txt}</h3>")
            continue

        # Form-specific: lines of mostly underscores or tabs become fill rows
        if is_form and re.match(r'^[\s_]{10,}$', text):
            lines.append('<p><span class="fill" style="min-width:400px"></span></p>')
            continue

        lines.append(f"<p>{txt}</p>")

    close_all_lists()

    # Append tables
    for table in tables:
        lines.append(render_table(table))

    return "\n".join(lines)


def build_html(doc_id: str, title: str, standard: str, body_html: str) -> str:
    prefix = doc_id.split("-")[0]
    if prefix == "POL":
        kind = "Policy"
    elif prefix == "PRO":
        kind = "Protocol"
    elif prefix == "FRM":
        kind = "Form"
    else:
        kind = "Document"

    authorised_by = infer_authorised_by(doc_id)
    standard_field = standard if standard else "AASM Standards for Accreditation"

    # Build clause tags for subtitle
    clause_tags_html = ""
    if standard:
        for clause in re.split(r',\s*', standard):
            clause = clause.strip()
            if clause:
                clause_tags_html += f'<span class="clause-tag">{html_module.escape(clause)}</span> '

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{html_module.escape(doc_id)} - {html_module.escape(title)}</title>
<style>
{CSS}
</style>
</head>
<body>
<div class="page">
  <header class="doc-header">
    <div class="kind">{html_module.escape(kind)}</div>
    <div class="title">{html_module.escape(title)}</div>
    <div class="subtitle">AASM Standards for Accreditation {clause_tags_html}</div>
    <div class="field"><strong>Document ID:</strong> {html_module.escape(doc_id)}</div>
    <div class="field"><strong>Revision:</strong> 1.0</div>
    <div class="field"><strong>Effective date:</strong> <span class="fill-sm"></span></div>
    <div class="field"><strong>Review date:</strong> <span class="fill-sm"></span></div>
    <div class="field"><strong>Authorised by:</strong> {html_module.escape(authorised_by)}</div>
    <div class="field"><strong>Standard:</strong> {html_module.escape(standard_field)}</div>
  </header>

{body_html}

  <div class="signoff">
    <div class="sign-cell"><strong>Prepared by</strong><br><br><span class="fill"></span><br>Name / Date</div>
    <div class="sign-cell"><strong>Reviewed by</strong><br><br><span class="fill"></span><br>Name / Date</div>
    <div class="sign-cell"><strong>Approved by</strong><br><br><span class="fill"></span><br>{html_module.escape(authorised_by)} / Date</div>
  </div>

  <div class="footer">{html_module.escape(doc_id)} v1.0 · AASM Standards for Accreditation · Nexus 360 Accreditation · Uncontrolled if printed</div>
</div>
</body>
</html>"""


def docx_to_html(docx_path: Path, doc_id: str, title: str, standard: str) -> str:
    is_form = doc_id.startswith("FRM")
    doc = docx.Document(str(docx_path))
    body_html = paragraphs_to_html(doc.paragraphs, doc.tables, is_form)
    return build_html(doc_id, title, standard, body_html)


def make_filename(doc_id: str, title: str) -> str:
    # Sanitize title: replace spaces/special chars with underscores
    safe_title = re.sub(r'[^\w\s-]', '', title)
    safe_title = re.sub(r'[\s/\\]+', '_', safe_title).strip('_')
    return f"{doc_id}_{safe_title}.html"


def find_docx_files(base_dir: Path):
    return list(base_dir.rglob("*.docx"))


def match_file(docx_path: Path, mapping):
    path_str = str(docx_path)
    for pattern, doc_id, standard, title in mapping:
        if re.search(pattern, path_str, re.IGNORECASE):
            return doc_id, standard, title
    return None, None, None


def main():
    base_dir = Path(r"C:\Users\wef.CMPHQ\Claude\NATAacredit\AASM Policies")
    out_dir = Path(r"C:\Users\wef.CMPHQ\Claude\NATAacredit\sop")
    out_dir.mkdir(parents=True, exist_ok=True)

    all_docx = find_docx_files(base_dir)
    print(f"Found {len(all_docx)} .docx files\n")

    successes = []
    failures = []
    skipped = []

    for docx_path in sorted(all_docx):
        doc_id, standard, title = match_file(docx_path, MAPPING)
        if not doc_id:
            skipped.append(str(docx_path))
            print(f"  SKIP (no mapping): {docx_path.name}")
            continue

        try:
            html_content = docx_to_html(docx_path, doc_id, title, standard)
            out_filename = make_filename(doc_id, title)
            out_path = out_dir / out_filename
            out_path.write_text(html_content, encoding="utf-8")
            successes.append(out_filename)
            print(f"  OK  {doc_id:15s}  → {out_filename}")
        except Exception as e:
            failures.append((str(docx_path), str(e)))
            print(f"  ERR {docx_path.name}: {e}")

    print(f"\n{'='*60}")
    print(f"Converted:  {len(successes)}")
    print(f"Skipped:    {len(skipped)}")
    print(f"Errors:     {len(failures)}")

    if skipped:
        print("\nSkipped files:")
        for s in skipped:
            print(f"  {s}")

    if failures:
        print("\nFailed files:")
        for path, err in failures:
            print(f"  {path}\n    Error: {err}")

    if successes:
        print("\nOutput files:")
        for f in sorted(successes):
            print(f"  {f}")


if __name__ == "__main__":
    main()
